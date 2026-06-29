#!/usr/bin/env python3
"""Add table borders + header-row shading to a pandoc-generated .docx."""
import sys
from docx import Document
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

path = sys.argv[1]
doc = Document(path)

def set_borders(tbl):
    tblPr = tbl.tblPr
    for ex in tblPr.findall(qn('w:tblBorders')):
        tblPr.remove(ex)
    b = OxmlElement('w:tblBorders')
    for edge in ('top','left','bottom','right','insideH','insideV'):
        e = OxmlElement(f'w:{edge}')
        e.set(qn('w:val'),'single'); e.set(qn('w:sz'),'4'); e.set(qn('w:space'),'0'); e.set(qn('w:color'),'808080')
        b.append(e)
    tblPr.append(b)

for t in doc.tables:
    set_borders(t._tbl)
    if t.rows:
        for cell in t.rows[0].cells:
            tcPr = cell._tc.get_or_add_tcPr()
            shd = OxmlElement('w:shd'); shd.set(qn('w:val'),'clear'); shd.set(qn('w:fill'),'E8EEF0')
            tcPr.append(shd)
            for p in cell.paragraphs:
                for r in p.runs: r.font.bold = True

doc.save(path)
print(f"post-processed {len(doc.tables)} tables")
